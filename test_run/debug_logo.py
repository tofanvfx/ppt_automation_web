import sys
import os
from pptx import Presentation
from docx import Document

# Create dummy docx
docx_path = 'test_run/dummy_logo.docx'
doc = Document()
doc.add_paragraph('[sst_content_page_01]')
doc.add_paragraph('Topic: T')
doc.save(docx_path)

# Mock generate_ppt locally with prints
from app.generators.lower_class_generator.docx_to_ppt import generate_ppt
import app.generators.lower_class_generator.docx_to_ppt as gen_module

# Add a wrapper for print
original_inject = gen_module.inject_logo if hasattr(gen_module, 'inject_logo') else None

def debug_ppt():
    pptx_path = 'test_run/debug_logo_out.pptx'
    template_path = 'app/generators/lower_class_generator/template.pptx'
    
    # We will read template
    prs = Presentation(template_path)
    logos = []
    for layout in prs.slide_layouts:
        if layout.name == "LAYOUT_logo":
            logos = list(layout.shapes)
            print("FOUND LAYOUT_logo shapes:", len(logos))
            for s in logos:
                print(" - Shape name:", getattr(s, 'name', 'N/A'), "Type:", s.shape_type)
            break
            
debug_ppt()

"""
Inspect DOCX bullets by creating a real Word-style multi-level list.
We use python-docx but also patch the numPr to set real ilvl values.
"""
import sys, os
sys.path.insert(0, os.getcwd())
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

docx_path = 'test_run/real_bullets.docx'
doc = Document()
doc.add_paragraph('[sst_content_page_01]')

# Add topic/subtopic
p = doc.add_paragraph('Topic: Establishment of the Delhi Sultanate')
p = doc.add_paragraph('Subtopic: Ruling Dynasties')

# Level 0 bullet - standard List Paragraph style
def add_list_para(doc, text, ilvl=0, numId=1, bold=False):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = None
    run = p.add_run(text)
    run.bold = bold

    # Set numPr with ilvl
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(numId))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    
    pPr = p._p.get_or_add_pPr()
    pPr.insert(0, numPr)
    return p

add_list_para(doc, 'Five successive Turkic-Afghan dynasties ruled:', ilvl=0)
add_list_para(doc, 'Mamluks (Slave Dynasty)', ilvl=1, bold=True)
add_list_para(doc, 'Khiljis', ilvl=1, bold=True)
add_list_para(doc, 'Tughlaqs', ilvl=1, bold=True)
add_list_para(doc, 'Sayyids', ilvl=1, bold=True)
add_list_para(doc, 'Lodis', ilvl=1, bold=True)

doc.save(docx_path)
print(f'Created: {docx_path}')

# Verify ilvl is set
doc2 = Document(docx_path)
w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for para in doc2.paragraphs:
    if not para.text.strip():
        continue
    numPr = para._element.find(f'.//{{{w_ns}}}numPr')
    ilvl_val = 0
    is_list = False
    if numPr is not None:
        is_list = True
        ilvl_el = numPr.find(f'{{{w_ns}}}ilvl')
        if ilvl_el is not None:
            ilvl_val = int(ilvl_el.get(f'{{{w_ns}}}val', '0'))
    print(f'  is_list={is_list} ilvl={ilvl_val} text={para.text[:40]!r}')
